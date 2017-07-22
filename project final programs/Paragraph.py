'''
Created on 10-Oct-2016

@author: Aishwarya-HP
'''
from docx import Document
from ai.liver import liver
from ai.tmpspeech import tmpspeech
class Paragraph:
    
    
    def liverP(self,fname):
        
        doc = Document(fname)
        docl=Document("liverp.docx")
        
        for pl in docl.paragraphs:
            pf=doc.add_paragraph(pl.text)
        doc.save(fname)
           
    def gallbladder(self,fname):
        
        doc = Document(fname)
        docl=Document("gallbadderp.docx")
        
        for pl in docl.paragraphs:
            pf=doc.add_paragraph(pl.text)
            
        doc.save(fname)
    
    def kidney(self,fname):
        doc = Document(fname)
        docl=Document("kidney.docx")
        
        for pl in docl.paragraphs:
            pf=doc.add_paragraph(pl.text)
            
        doc.save(fname)
    
    def uterus(self,fname):
        doc = Document(fname)
        docl=Document("uterus.docx")
        
        for pl in docl.paragraphs:
            pf=doc.add_paragraph(pl.text)
            
        doc.save(fname)
            
    def impression(self,fname):
            doc = Document(fname)
            impstr=""
            finalstr=""
            doc.add_paragraph('impression')
            while(impstr!='ok'):
                print("tell")
                
                
                obj=tmpspeech()
                impstr=str(obj.speech())
                print(impstr)
                if(impstr=="ok"):
                    break
                finalstr=finalstr+impstr
            print(finalstr)
            l=liver()    
            doc.add_paragraph(finalstr)
            doc.save(fname)
            l.replace('full stop','.',fname)
            l.replace('into','x',fname)
            l.replace('by','x',fname)
            doc.save(fname)
            
    def general(self,fname):
        doc=Document(fname)
        impstr="empty"
        finalstr=""
        while(impstr!='ok'):
            print("tell  say ok to stop")
            
            obj=tmpspeech()
            impstr=str(obj.speech())
            print(impstr)
           
            if(impstr=="ok"):
                break
            if(impstr=='impression'):
                doc.add_paragraph("Impression :")
                doc.save(fname)
                continue
                
            print('say final to finalise the spoken statement')
            str1=str(obj.speech())
            
            if(str1=='final'):
                print("statement added")
                finalstr=finalstr+" "+impstr
            else:
                continue
            
        print(finalstr)
        l=liver()
        doc.add_paragraph(finalstr)
        doc.save(fname)
        l.replace('full stop','.',fname)
        l.replace('into','x',fname)
        l.replace('by','x',fname)
        doc.save(fname)
        
        
        
        
