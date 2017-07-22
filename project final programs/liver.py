'''
Created on 10-Sep-2016

@author: Aishwarya-HP
'''
import os
from docx import Document
from ai.tmpspeech import tmpspeech


class liver:
    def replace(self,old,new,fname):
        
        doc = Document(fname)
        for p in doc.paragraphs:
            if old in p.text:
                text = p.text.replace(old,new)
                style = p.style
                p.text = text
                p.style = style
            
        doc.save(fname)
    
    def p1(self,fname,pname):
        self.replace('p1',pname,fname)
    
    def id2(self,fname,pid):
        self.replace('id2',pid,fname)   
        
    def l1(self,fname):
        para=self.clears2()
        self.replace('L1',para,fname)
        
    def l2(self,fname):
        para=self.clears2()
        self.replace('L2',para,fname)
    
    def l3(self,fname):
        para=self.clears2()
        self.replace('L3',para,fname)
    
    def l4(self,fname):
        para=self.clears2()
        self.replace('L4',para,fname)
    
    def l5(self,fname):
        para=self.clears2()
        self.replace('L5',para,fname)
    
        
    def clears(self):
        while(1):
            print("tell")
            obj=tmpspeech()
            para=str(obj.speech())
            print(para)
            print("speak final to stop")
            tmp=obj.speech()
            print(tmp)
            if(tmp=="final"):
                break
        return para
    
    def clears2(self):
        while(1):
            print("tell")
            obj=tmpspeech()
            para=str(obj.speech())
            print(para)
            if(len(para)==0):
                continue
            else:
                break
            
        return para
    
            
    def k1(self,fname):
        para=self.clears()
        self.replace('k1',para,fname)
        self.replace('into',' x ',fname)
        self.replace('by',' x ',fname)
    
    def k2(self,fname):
        para=self.clears()
        self.replace('k2',para,fname)
        self.replace('into',' x ',fname)
        self.replace('by',' x ',fname)
    
    def stone(self,fname):
        pa="empty"
        while(1):
            print("tell stone present or not") 
            obj=tmpspeech()
            pa=str(obj.speech())
            print(pa)
            if(pa=="present" or pa=="absent"):
                break
            else:
                continue
        if(pa=='present'):
            self.replace('STONE','size of kidney stone : K3',fname)
            print("tell size of kidney stone")
            para=self.clears()
            self.replace('K3',para,fname)
            self.replace('into',' x ',fname)
            self.replace('by',' x ',fname)
        else:
            self.replace('STONE','no echoreflective calculus found.',fname)
            
    def hyper(self,fname):
        pa="empty"
        while(1):
            print("is RENOVASCULAR_HYPERTENSION found say yes or no") 
            obj=tmpspeech()
            pa=str(obj.speech())
            print(pa)
            if(pa=="yes" or pa=="no"):
                break
            else:
                continue
        if(pa=="yes"):
            self.replace('RENOVASCULAR_HYPERTENSION','Renovascular hypertension found.',fname)
        else:
            self.replace('RENOVASCULAR_HYPERTENSION','Renovascular hypertension not found.',fname)
    
    def display(self,filename):
        self.replace('full stop','.',filename)
        os.startfile(filename)
    
    def filecopy(self,fname):    
        source_document = Document('report.docx')
        target_document = Document(fname)

        for paragraph in source_document.paragraphs:
            text = paragraph.text
            target_document.add_paragraph(text)
    
        target_document.save(fname)    
        
    def createf(self,fname):
        document = Document()
        document.save(fname)